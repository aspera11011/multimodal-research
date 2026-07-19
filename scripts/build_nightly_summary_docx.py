#!/usr/bin/env python3
"""Build a plain-language nightly research summary for an internal supervisor."""

from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "reports" / "2026-07-18_19_nightly_research_summary.docx"
ASSETS = ROOT / "reports" / "assets"
QA_IMAGE = (
    ROOT
    / "results"
    / "qualitative"
    / "material_constancy_albedo_v1"
    / "smoke"
    / "full_comparison.jpg"
)
SKILL_SCRIPTS = Path(
    r"C:\Users\asper\.codex\plugins\cache\openai-primary-runtime\documents\26.715.12143\skills\documents\scripts"
)
sys.path.insert(0, str(SKILL_SCRIPTS))
from table_geometry import apply_table_geometry, column_widths_from_weights, section_content_width_dxa  # noqa: E402


NAVY = "17365D"
BLUE = "2E75B6"
TEAL = "2F7F7B"
GOLD = "C58A22"
INK = "263238"
MUTED = "66727C"
LIGHT_BLUE = "EAF2F8"
LIGHT_TEAL = "E9F4F2"
LIGHT_GOLD = "FFF5DD"
LIGHT_GRAY = "F3F5F7"
WHITE = "FFFFFF"
RED = "A23B3B"
FONT = "Microsoft YaHei"


def set_run_font(run, size=10.5, bold=False, color=INK, italic=False):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    return run


def set_cell_fill(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_border(cell, **edges):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge_name, attrs in edges.items():
        tag = f"w:{edge_name}"
        edge = borders.find(qn(tag))
        if edge is None:
            edge = OxmlElement(tag)
            borders.append(edge)
        for key, value in attrs.items():
            edge.set(qn(f"w:{key}"), str(value))


def no_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant = OxmlElement("w:cantSplit")
    tr_pr.append(cant)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, 8.5, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)
    run = paragraph.add_run(" 页")
    set_run_font(run, 8.5, color=MUTED)


def add_para(doc, text="", size=10.5, bold=False, color=INK, after=6, before=0, align=None, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.16
    if align is not None:
        p.alignment = align
    set_run_font(p.add_run(text), size=size, bold=bold, color=color, italic=italic)
    return p


def add_heading(doc, text, level=1):
    sizes = {1: 16, 2: 12.5, 3: 11}
    colors = {1: NAVY, 2: BLUE, 3: TEAL}
    before = {1: 2, 2: 8, 3: 6}
    after = {1: 7, 2: 5, 3: 4}
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(before[level])
    p.paragraph_format.space_after = Pt(after[level])
    set_run_font(p.add_run(text), sizes[level], True, colors[level])
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.65)
    p.paragraph_format.first_line_indent = Cm(-0.3)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.12
    if bold_prefix and text.startswith(bold_prefix):
        set_run_font(p.add_run(bold_prefix), 10.3, True, INK)
        set_run_font(p.add_run(text[len(bold_prefix) :]), 10.3, False, INK)
    else:
        set_run_font(p.add_run(text), 10.3, False, INK)
    return p


def add_callout(doc, title, body, fill=LIGHT_BLUE, accent=BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.18)
    p.paragraph_format.right_indent = Cm(0.08)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:color"), accent)
    left.set(qn("w:space"), "7")
    borders.append(left)
    p_pr.append(borders)
    set_run_font(p.add_run(title), 11, True, accent)
    p.add_run().add_break()
    set_run_font(p.add_run(body), 10.3, False, INK)
    return p


def add_table(doc, headers, rows, weights, header_fill=NAVY, font_size=9.4):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    repeat_header(hdr)
    no_split(hdr)
    for idx, text in enumerate(headers):
        cell = hdr.cells[idx]
        set_cell_fill(cell, header_fill)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(p.add_run(text), font_size, True, WHITE)
    for row_idx, values in enumerate(rows):
        row = table.add_row()
        no_split(row)
        if row_idx % 2:
            for cell in row.cells:
                set_cell_fill(cell, LIGHT_GRAY)
        for idx, value in enumerate(values):
            cell = row.cells[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_run_font(p.add_run(str(value)), font_size, False, INK)
    widths = column_widths_from_weights(weights, section_content_width_dxa(doc.sections[-1]))
    apply_table_geometry(table, widths)
    add_para(doc, "", after=2)
    return table


def add_picture_with_alt(doc, path, width, alt, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run()
    inline = run.add_picture(str(path), width=width)
    doc_pr = inline._inline.docPr
    doc_pr.set("descr", alt)
    cap = add_para(doc, caption, size=8.8, color=MUTED, after=6, align=WD_ALIGN_PARAGRAPH.CENTER)
    cap.paragraph_format.keep_with_next = False


def page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def make_chart(path):
    labels = ["RGB", "Albedo", "RGB+Albedo", "RGB+Text"]
    qwen_acc = [60.30, 36.36, 61.21, 36.67]
    intern_acc = [53.03, 29.39, 51.21, 25.15]
    qwen_flip = [57.58, 71.21, 50.00, 36.36]
    intern_flip = [66.67, 77.27, 65.15, 36.36]
    canvas = Image.new("RGB", (1800, 560), "white")
    draw = ImageDraw.Draw(canvas)
    font_path = Path(r"C:\Windows\Fonts\msyh.ttc")
    bold_path = Path(r"C:\Windows\Fonts\msyhbd.ttc")
    font = lambda size, bold=False: ImageFont.truetype(str(bold_path if bold and bold_path.exists() else font_path), size)
    draw.text((900, 24), "Same 330 samples under different evidence interfaces", fill="#263238", font=font(30, True), anchor="ma")
    panels = [
        (70, "Accuracy (%)", qwen_acc, intern_acc),
        (930, "Cross-light flip rate (%)", qwen_flip, intern_flip),
    ]
    for left, title, qvals, ivals in panels:
        top, bottom, width = 105, 455, 760
        draw.text((left + width / 2, 78), title, fill="#17365D", font=font(24, True), anchor="ma")
        draw.line((left, bottom, left + width, bottom), fill="#87929B", width=2)
        for tick in [0, 20, 40, 60, 80]:
            y = bottom - int((tick / 85) * (bottom - top))
            draw.line((left, y, left + width, y), fill="#E2E7EB", width=1)
            draw.text((left - 10, y), str(tick), fill="#66727C", font=font(16), anchor="rm")
        group_w = width / 4
        bar_w = 58
        for idx, label in enumerate(labels):
            center = left + group_w * (idx + 0.5)
            for offset, value, color in [(-bar_w / 2, qvals[idx], "#2E75B6"), (bar_w / 2, ivals[idx], "#2F7F7B")]:
                x0 = center + offset - bar_w / 2
                x1 = x0 + bar_w
                y = bottom - (value / 85) * (bottom - top)
                draw.rectangle((x0, y, x1, bottom), fill=color)
                draw.text(((x0 + x1) / 2, y - 5), f"{value:.0f}", fill="#263238", font=font(15, True), anchor="ms")
            draw.text((center, bottom + 18), label, fill="#263238", font=font(15), anchor="ma")
    draw.rectangle((720, 500, 750, 525), fill="#2E75B6")
    draw.text((760, 512), "Qwen", fill="#263238", font=font(17), anchor="lm")
    draw.rectangle((890, 500, 920, 525), fill="#2F7F7B")
    draw.text((930, 512), "InternVL", fill="#263238", font=font(17), anchor="lm")
    canvas.save(path, quality=95)


def build():
    ASSETS.mkdir(parents=True, exist_ok=True)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    chart = ASSETS / "nightly_summary_metrics.png"
    make_chart(chart)

    doc = Document()
    section = doc.sections[0]
    # Named override to standard_business_brief: A4 and Chinese office typography.
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(0.85)
    section.footer_distance = Cm(0.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.16
    for style_name in ["Heading 1", "Heading 2", "Heading 3", "List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run_font(header.add_run("多模态物理证据研究 | 阶段汇报"), 8.5, True, MUTED)
    footer = section.footer.paragraphs[0]
    add_page_number(footer)

    # Page 1
    add_para(doc, "阶段研究汇报", 10, True, GOLD, after=8)
    add_para(doc, "今晚的工作与核心 finding", 25, True, NAVY, after=5)
    add_para(doc, "从支撑关系排查，到材质恒常性三步实验", 13, False, BLUE, after=13)
    add_para(doc, "2026年7月18日晚 - 7月19日凌晨  |  内部阶段总结", 9.3, False, MUTED, after=16)
    add_callout(
        doc,
        "一句话结论",
        "我们找到了一个真实且可重复的问题：同一块材质只是换了灯光，VLM就经常改变答案。但官方albedo模型虽然能把图像变得更稳定，冻结VLM却不会直接利用它。因此“零训练直接加提示”这条路线暂停。",
        LIGHT_BLUE,
        BLUE,
    )
    add_heading(doc, "今晚最重要的三个数字", 2)
    add_table(
        doc,
        ["跨光照翻转", "两模型共同翻转", "Albedo像素波动"],
        [["Qwen 57.58%\nInternVL 66.67%", "34/66 个区域\n51.52%", "0.199 -> 0.083\n65/66区域改善"]],
        [1, 1, 1],
        header_fill=TEAL,
        font_size=10,
    )
    add_heading(doc, "这个结果意味着什么", 2)
    add_bullet(doc, "问题是存在的，而且不是某一个模型的偶然现象。")
    add_bullet(doc, "物理证据本身可以更稳定，但“怎么让VLM读懂”是真正的难点。")
    add_bullet(doc, "今晚没有得到一个可以直接写论文的方法，但成功定位了问题和无效接口，避免继续盲目投入。")

    page_break(doc)
    # Page 2
    add_heading(doc, "1. 今晚实际做了哪些工作", 1)
    add_para(doc, "整个过程不是直接跑一次模型，而是按“数据是否可用 - 问题是否存在 - 方法能否修复”逐层排查。")
    add_table(
        doc,
        ["阶段", "实际完成的工作", "得到的判断"],
        [
            ["服务器与留痕", "完成SSH密钥、校园网/GitHub连通、个人工作区、GPU 3使用与本地-远程同步。", "后续实验可复现，不触碰服务器其他人目录。"],
            ["Support关系pilot", "下载并校验NYUv2；从1,449帧中构造167个候选，人工检查14个。", "仅2个是支撑，11个不是，1个不确定；数据不适合规模化做真值。"],
            ["方向查重与转向", "核验InstaFormer等相关工作，排除遮挡文本提示的高重合路线；审计材质恒常性数据。", "选择“光照变化下的VLM材质稳定性”作为低成本验证项。"],
            ["三步材质实验", "从23区域pilot扩到66区域/330样本，复核两个VLM，再运行Marigold IID和三种干预接口。", "确认RGB-only故障，但零训练albedo干预未通过。"],
        ],
        [1.25, 3.4, 2.15],
        font_size=8.8,
    )
    add_heading(doc, "为什么Support方向要停", 2)
    add_para(doc, "深度和实例框可以帮我们找到“上下靠得很近”的物体，但这不等于它们真的直接接触和支撑。隔板、柜子、大物体mask等都会产生高分误报。继续人工标注成本很高，而正例又很少，所以及时停损比继续堆数据更合理。")
    add_callout(doc, "这部分不是白做", "它让我们在只人工审核14个样本时就发现数据根本限制，避免了将数百个几何候选误当成真实支撑关系。", LIGHT_GOLD, GOLD)

    page_break(doc)
    # Page 3
    add_heading(doc, "2. 材质恒常性：问题是真的", 1)
    add_para(doc, "材质恒常性可以简单理解为：同一个物体换了灯光，我们仍应判断它是同一种材质。实验中的物体、位置和材质都不变，只改变真实灯光。")
    add_table(
        doc,
        ["实验步骤", "数据规模", "结果", "判断"],
        [
            ["第一步小pilot", "10场景\n23区域\n115张图", "Qwen翻转 47.83%", "问题初步存在"],
            ["第二步扩展复核", "30场景\n66区域\n330张图", "Qwen 57.58%\nInternVL 66.67%", "两个模型都存在"],
            ["两模型共同失败", "同一批66区域", "34/66同时翻转", "不是Qwen的偶然问题"],
        ],
        [1.45, 1.35, 2.1, 1.9],
        font_size=9.1,
    )
    add_picture_with_alt(
        doc,
        chart,
        Inches(6.25),
        "柱状图：Qwen与InternVL在RGB、albedo、RGB加albedo和RGB加属性文本条件下的准确率和翻转率。",
        "图1  同一330个样本在不同证据接口下的表现（翻转率越低越好，但必须同时看准确率）",
    )
    add_heading(doc, "如何理解“翻转”", 2)
    add_para(doc, "如果同一块布在5种灯光下，模型有4次说“布料”、1次说“瓷砖”，这个区域就算发生过翻转。这个指标不只看某一张图答对没有，而是看模型对同一材质是否稳定。")
    add_callout(doc, "已验证的 finding", "在有明显照明差异的压力测试中，冻结VLM的材质答案会大量随灯光变化；且两个不同系列的2B模型在容易翻转的区域上高度重合。", LIGHT_TEAL, TEAL)

    page_break(doc)
    # Page 4
    add_heading(doc, "3. Albedo干预：图像稳定了，VLM却没有变好", 1)
    add_para(doc, "Albedo可以理解为尽量去掉灯光影响后的表面固有外观。我们使用官方Marigold IID Appearance v1.1，先处理247张完整场景-光照图，再按原坐标裁剪，避免直接对小图分解带来假结果。")
    if QA_IMAGE.exists():
        add_picture_with_alt(
            doc,
            QA_IMAGE,
            Inches(6.25),
            "左侧为原RGB室内场景，右侧为Marigold生成的albedo；右侧阴影和明暗变化明显减少，但存在色彩偏移和局部误差。",
            "图2  同一张室内场景：左为RGB，右为Marigold albedo",
        )
    add_table(
        doc,
        ["输入方式", "Qwen准确率 / 翻转", "InternVL准确率 / 翻转", "结论"],
        [
            ["RGB基线", "60.30% / 57.58%", "53.03% / 66.67%", "对照组"],
            ["Albedo-only", "36.36% / 71.21%", "29.39% / 77.27%", "更差"],
            ["RGB + Albedo", "61.21% / 50.00%", "51.21% / 65.15%", "无稳定收益"],
            ["RGB + 物理数值", "36.67% / 36.36%", "25.15% / 36.36%", "伪一致性"],
        ],
        [1.4, 1.9, 1.9, 1.6],
        font_size=8.8,
    )
    add_heading(doc, "最关键的反例", 2)
    add_para(doc, "把Marigold预测的粗糙度和金属度写成文本后，两个模型确实“更稳定”了，但Qwen有227/330次、InternVL有269/330次直接回答“metal”。这种稳定是因为它们反复答同一个错误类别，不是真正理解了材质。")
    add_callout(doc, "第三步决策：No-Go", "“直接给albedo图”和“直接给物理数值文本”都没有在两个模型上同时提高准确率和稳定性。按实验前设定的停损标准，这条零训练路线暂停。", "FCECEC", RED)

    page_break(doc)
    # Page 5
    add_heading(doc, "4. 今晚的最终 finding 与建议", 1)
    add_heading(doc, "可以较有把握地说", 2)
    add_bullet(doc, "冻结VLM的材质判断对真实灯光变化很敏感，该现象已在Qwen和InternVL上重复。")
    add_bullet(doc, "Marigold生成的albedo确实更稳定，说明“消除照明干扰”这个物理出发点没有问题。")
    add_bullet(doc, "真正的瓶颈是VLM不会直接使用本征图或物理数值，存在明显的“物理证据-模型接口”鸿沟。")
    add_heading(doc, "现在还不能说", 2)
    add_bullet(doc, "还不能说我们已经有了一个可发表的新方法。")
    add_bullet(doc, "还不能把压力测试中的57%-67%翻转率，理解为日常随机灯光下的自然发生概率。")
    add_bullet(doc, "只测了2B级模型，且clear plastic、marble等类别仍较少，还不能外推到所有VLM。")
    add_heading(doc, "建议怎么处理这条线", 2)
    add_callout(
        doc,
        "建议：归档当前零训练方案，不再继续堆实验",
        "如果组里希望继续，研究问题必须升级为“如何学习albedo/材质属性与VLM之间的对齐接口”，例如轻量adapter或微调。这会明显增加数据、训练和查重成本，不再是原先预期的低成本DOP式课题。",
        LIGHT_GOLD,
        GOLD,
    )
    add_heading(doc, "今晚留下的可复现材料", 2)
    add_table(
        doc,
        ["材料", "已保存内容", "位置"],
        [
            ["数据与manifest", "66区域、330样本、两个本征证据manifest", "local/multimodal-research/experiments"],
            ["模型预测", "RGB、albedo、RGB+albedo、物理文本的逐样本输出", "local/multimodal-research/results"],
            ["脚本和配置", "数据构建、Qwen/InternVL推理、Marigold、bootstrap分析", "local/multimodal-research/scripts / configs"],
            ["运行日志", "下载、环境、全量推理与结果哈希", "local/multimodal-research/experiments/logs"],
        ],
        [1.25, 3.2, 2.35],
        font_size=8.4,
    )
    add_para(doc, "备注：这份文档是内部阶段汇报。所有结果都保留了配置、逐样本预测、日志和本地-远程哈希；数据和权重不进Git。", 8.6, False, MUTED, after=2)

    core = doc.core_properties
    core.title = "今晚的工作与核心finding"
    core.subject = "多模态物理证据研究阶段汇报"
    core.author = "Research Project"
    core.keywords = "VLM, material constancy, albedo, nightly summary"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
